"""
export_report.py — Automated Report Generation Module
Public Health Surveillance Data Pipeline
Author: Aly Drame, MD, MPH, MBA

Generates visualizations and an automated Word document
surveillance summary report.
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as mtick
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

plt.style.use('seaborn-v0_8-whitegrid')
sns.set_palette('Set2')


def plot_epidemic_curve(weekly_df: pd.DataFrame,
                         diseases: list = None,
                         output_path: str = 'outputs/epidemic_curve.png'
                         ) -> None:
    """
    Generate epidemic curve line chart by disease and week.
    """
    df = weekly_df.copy()
    if diseases:
        df = df[df['disease_name'].isin(diseases)]

    top_diseases = (df.groupby('disease_name')['case_count']
                    .sum().nlargest(5).index.tolist())
    df = df[df['disease_name'].isin(top_diseases)]

    fig, ax = plt.subplots(figsize=(14, 5))
    colors = sns.color_palette('Set2', len(top_diseases))

    for idx, disease in enumerate(top_diseases):
        sub = df[df['disease_name'] == disease].sort_values(
            ['epi_year','epi_week_num'])
        ax.plot(range(len(sub)), sub['case_count'],
                label=disease, color=colors[idx],
                linewidth=2, marker='o', markersize=3)

    ax.set_xlabel('Epidemiologic Week', fontsize=12)
    ax.set_ylabel('Case Count', fontsize=12)
    ax.set_title('Weekly Reported Cases by Disease\n'
                 'Synthetic Surveillance Data — Portfolio Demonstration',
                 fontsize=13, fontweight='bold')
    ax.legend(loc='upper right', fontsize=9)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f"[REPORT] Epidemic curve saved → {output_path}")


def plot_completeness_heatmap(miss_df: pd.DataFrame,
                               output_path: str = 'outputs/completeness_heatmap.png'
                               ) -> None:
    """
    Heatmap of data completeness by field.
    Red = low completeness, green = high completeness.
    """
    plot_data = miss_df[miss_df['pct_missing'] > 0].copy()
    if plot_data.empty:
        print("[REPORT] No missing data to plot in heatmap")
        return

    pivot = plot_data.set_index('field')['completeness_pct'].to_frame()

    fig, ax = plt.subplots(figsize=(4, max(4, len(pivot) * 0.4)))
    sns.heatmap(pivot, annot=True, fmt='.1f', cmap='RdYlGn',
                vmin=70, vmax=100, ax=ax, cbar_kws={'label': '% Complete'},
                linewidths=0.5)
    ax.set_title('Data Completeness by Field\n'
                 'Synthetic Surveillance Data', fontsize=12,
                 fontweight='bold')
    ax.set_xlabel('')
    ax.set_ylabel('')
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f"[REPORT] Completeness heatmap saved → {output_path}")


def plot_disease_distribution(df: pd.DataFrame,
                               output_path: str = 'outputs/disease_distribution.png'
                               ) -> None:
    """Horizontal bar chart of case counts by disease."""
    dist = (df.groupby('disease_name')['case_id']
            .count().reset_index()
            .rename(columns={'case_id':'case_count'})
            .sort_values('case_count', ascending=True))

    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.barh(dist['disease_name'], dist['case_count'],
                   color=sns.color_palette('Set2', len(dist)))
    for bar, val in zip(bars, dist['case_count']):
        ax.text(val + 10, bar.get_y() + bar.get_height() / 2,
                f'{val:,}', va='center', fontsize=9)
    ax.set_xlabel('Case Count', fontsize=12)
    ax.set_title('Reported Cases by Disease\n'
                 'Synthetic Surveillance Data — Portfolio Demonstration',
                 fontsize=13, fontweight='bold')
    ax.set_xlim(0, dist['case_count'].max() * 1.15)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f"[REPORT] Disease distribution saved → {output_path}")


def create_word_report(weekly_df: pd.DataFrame,
                        miss_df: pd.DataFrame,
                        jur_df: pd.DataFrame,
                        epidemic_curve_path: str,
                        output_path: str = 'outputs/weekly_surveillance_report.docx'
                        ) -> None:
    """
    Generate a formatted Word document surveillance summary report.
    Mirrors the weekly situation reports produced in CDC programs.
    """
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────
    title = doc.add_heading(
        'Weekly Public Health Surveillance Summary Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        f'Report Date: {datetime.now().strftime("%B %d, %Y")}  |  '
        f'Data: Synthetic — Portfolio Demonstration')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph()

    # ── Section 1: Executive Summary ───────────────────────────────
    doc.add_heading('1. Executive Summary', level=1)

    total_cases = len(weekly_df['case_count'].dropna()) \
                  if 'case_count' in weekly_df.columns else 'N/A'
    n_diseases  = weekly_df['disease_name'].nunique() \
                  if 'disease_name' in weekly_df.columns else 'N/A'

    doc.add_paragraph(
        f'This report summarizes synthetic surveillance data covering '
        f'{n_diseases} notifiable diseases across 15 jurisdictions '
        f'for the period 2022 through 2024. All data are synthetic and '
        f'generated for portfolio demonstration purposes.')

    doc.add_paragraph()

    # ── Section 2: Data Quality ─────────────────────────────────────
    doc.add_heading('2. Data Quality Summary', level=1)
    doc.add_paragraph(
        'Fields with completeness below 95% are flagged. High missing '
        'rates in race/ethnicity and onset_date are consistent with known '
        'surveillance system limitations.')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Field','Total Records',
                            'Missing Count','% Complete']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    top_missing = miss_df.head(10)
    for _, row in top_missing.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['field'])
        cells[1].text = f"{int(row['total_records']):,}"
        cells[2].text = f"{int(row['missing_count']):,}"
        cells[3].text = f"{row['completeness_pct']:.1f}%"

    doc.add_paragraph()

    # ── Section 3: Epidemic Curve ───────────────────────────────────
    doc.add_heading('3. Weekly Case Counts — Epidemic Curve', level=1)
    if os.path.exists(epidemic_curve_path):
        doc.add_picture(epidemic_curve_path, width=Inches(6.0))
    else:
        doc.add_paragraph('[Epidemic curve chart — see outputs/epidemic_curve.png]')

    doc.add_paragraph()

    # ── Section 4: Case Summary Table ──────────────────────────────
    doc.add_heading('4. Disease Summary', level=1)
    if not weekly_df.empty and 'disease_name' in weekly_df.columns:
        summary = (
            weekly_df.groupby('disease_name')
            .agg(
                total_cases    = ('case_count', 'sum'),
                hospitalizations = ('hospitalized_count', 'sum'),
                deaths         = ('death_count', 'sum'),
            )
            .reset_index()
            .sort_values('total_cases', ascending=False)
        )
        table2 = doc.add_table(rows=1, cols=4)
        table2.style = 'Light Shading Accent 1'
        h2 = table2.rows[0].cells
        for i, h in enumerate(['Disease','Total Cases',
                                'Hospitalizations','Deaths']):
            h2[i].text = h
            h2[i].paragraphs[0].runs[0].bold = True
        for _, row in summary.iterrows():
            cells = table2.add_row().cells
            cells[0].text = str(row['disease_name'])
            cells[1].text = f"{int(row['total_cases']):,}"
            cells[2].text = f"{int(row['hospitalizations']):,}"
            cells[3].text = f"{int(row['deaths']):,}"

    doc.add_paragraph()

    # ── Footer ──────────────────────────────────────────────────────
    footer = doc.add_paragraph(
        'Note: All data are synthetic and generated for portfolio '
        'demonstration purposes only. This report format mirrors CDC '
        'surveillance program reporting structure. No real surveillance '
        'data were used.')
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)

    doc.save(output_path)
    print(f"[REPORT] Word report saved → {output_path}")
